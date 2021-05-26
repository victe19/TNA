import os
import io

from django.http import HttpResponse
from django.shortcuts import render, redirect
from django.views.generic import TemplateView

from google.cloud import vision
import docx
from wsgiref.util import FileWrapper

from google.oauth2 import service_account
from google.cloud import speech



from .forms import FileForm
from .models import File



class Home(TemplateView):
    template_name = 'home.html'

def about(request):
    return render(request,"about.html")

# Handle the file list view
def file_list(request):
    files = File.objects.all()
    return render(request, 'file_list.html', {
        'files': files
    })


# Handle the upload of a new model class (file)
def upload_file(request):
    if request.method == 'POST':
        form = FileForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            return redirect('file_list')
    else:
        form = FileForm()
    return render(request, 'upload_file.html', {
        'form': form
    })


# Handler for file deleting
def delete_file(request, pk):
    if request.method == 'POST':
        file = File.objects.get(pk=pk)
        file.delete()
    return redirect('file_list')


def start(request, pk):
    if request.method == 'GET':
        # Collect the file
        ressource = File.objects.get(pk=pk)

        filename, file_extension = os.path.splitext(str(ressource.file))
        credentials = service_account.Credentials.from_service_account_file("media/angelic-cat-312411-d70e72cccae5.json")

        if file_extension == ".wav":

            response = transcribe_audio(ressource.file,credentials)
            print(response)


        #imatge
        else:
            response = transcribe_image(ressource.file,credentials)
            if type(response) == "str":
                return render(request,"file_list.html",{"error":response})

        file_path = createDocx(response, ressource, file_extension)
        request.session["path"] = file_path

        return redirect("save")

    return render(request, "results.html", {'file': transcribe_audio})


def transcribe_audio(audio_ressource,credentials):

    # Instantiates a client
    client = speech.SpeechClient(credentials=credentials)

    # # resource at google cloud storage else it's local
    # if audio_ressource.startswith("gs://"):
    #     # gsc_uri = "gs://cloud-samples-data/speech/brooklyn_bridge.raw"
    #     print(audio_ressource)
    #     audio = speech.RecognitionAudio(uri=audio_ressource)

    name = str(audio_ressource.file)
    print(name)
    with io.open(name, 'rb') as audio_file:
        content = audio_file.read()
        audio = speech.RecognitionAudio(content=content)


    config = speech.RecognitionConfig(
        encoding=speech.RecognitionConfig.AudioEncoding.LINEAR16,
        sample_rate_hertz=44100,
        language_code='es-ES',  # todo es-ES
    )

    response = client.recognize(config=config, audio=audio)

    return response

def createDocx(response,resource,type):


    if type == ".wav":

        for i, result in enumerate(response.results):
            alternative = result.alternatives[0]
            print("-" * 20)
            print("First alternative of result {}".format(i))
            print("Transcript: {}".format(alternative.transcript))

        text = alternative.transcript
    else:
        text = response.text_annotations
        text = text[0].description

    words = text.split()

    print(" ----- Nou registre d'apunts ----- \n")
    # time.sleep(1)
    print("Emplena Informació: \n")
    # time.sleep(2)
    paraulesLinia = 9999


    ass = resource.subject
    prof = resource.prof

    dia = resource.day
    mes = resource.month

    newDoc = docx.Document()
    style = newDoc.styles['Normal']
    font = style.font
    font.name = 'Arial'

    newDoc.add_paragraph(str(dia) + '/' + str(mes) + "     " + prof)
    newDoc.add_heading(ass, 0)  # titol

    for i, item in enumerate(words):
        if (i % paraulesLinia == 0):
            p = newDoc.add_paragraph(item)
        else:
            p.add_run(" ")
            p.add_run(item)

    file_path = "media/Transcribed_File.docx"
    newDoc.save(file_path)

    print("Success!: Pujat fitxer al directori apunts/" )

    return file_path

def save(request):
    try:
        file_path = request.session["path"]
        print(file_path)
        wrapper = FileWrapper(open(file_path, 'rb'))
        response = HttpResponse(wrapper, content_type='application/force-download')
        response['Content-Disposition'] = 'inline; filename=' + os.path.basename(file_path)
        print("done")
        return response
    except Exception as e:
        print("exceptioN!!!")
        return None

def transcribe_image(image_resource,credentials):

    client = vision.ImageAnnotatorClient(credentials=credentials)

    name = str(image_resource.file)
    with io.open(name, "rb") as image_file:
        content = image_file.read()
        image = vision.Image(content=content)


    response = client.document_text_detection(image=image)

    if response.error.message:
        return "Image could not be converted to text!"

    return response


